import openai
import json
import os
from docx import Document

# Establece tu clave de API de OpenAI
openai.api_key = 'TU_CLAVE_DE_API_AQUÍ'  # Reemplaza con tu clave real

def generar_seccion_resultados_con_citas(metricas):
    prompt = (
        "Escribe la sección de 'Resultados' para un artículo científico sobre reidentificación por silueta. "
        "Incluye al menos 15 citas en formato IEEE dentro de esta sección. "
        "El sistema utiliza técnicas de Visión por Computadora y algoritmos de Machine Learning. "
        "A partir de las siguientes métricas de evaluación del modelo, genera la sección incluyendo la definición de cada métrica, su uso e interpretación. "
        "no Omitas nombres de personas y asegura que las citas estén correctamente referenciadas en formato IEEE.\n\n"
        f"Métricas de evaluación:\n{json.dumps(metricas, indent=4, ensure_ascii=False)}\n\n"
    )
    
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Eres un asistente especializado en redacción científica en el campo de la visión por computadora y aprendizaje automático."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2000  # Ajusta según sea necesario
    )
    resultados = response['choices'][0]['message']['content']
    return resultados

def main():
    # Ruta al archivo JSON de entrada
    input_json_path = os.path.join('data', 'input', 'resultados_metricas.json')
    
    # Verificar si el archivo JSON existe
    if not os.path.exists(input_json_path):
        print(f"El archivo JSON no existe: {input_json_path}")
        return
    
    # Leer el JSON 'resultados_metricas'
    with open(input_json_path, 'r', encoding='utf-8') as json_file:
        metricas = json.load(json_file)
    
    # Generar la sección de Resultados con citas
    resultados = generar_seccion_resultados_con_citas(metricas)
    
    # Crear el documento de Word y agregar el contenido
    document = Document()
    document.add_heading('Resultados', level=1)
    document.add_paragraph(resultados)
    
    # Crear la carpeta de salida si no existe
    output_dir = os.path.join('data', 'output')
    os.makedirs(output_dir, exist_ok=True)
    
    # Ruta para guardar el documento de Word
    output_docx_path = os.path.join(output_dir, 'RESULTADOS.docx')
    
    document.save(output_docx_path)
    
    print(f"Documento 'RESULTADOS.docx' generado y guardado en {output_docx_path}.")

if __name__ == '__main__':
    main()

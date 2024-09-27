import openai
import json
import os
from docx import Document

# Establece tu clave de API de OpenAI
openai.api_key = 'TU_CLAVE_DE_API_AQUÍ'  # Reemplaza con tu clave real

def generar_introduccion_con_citas(articulos):
    # Crear el contexto combinando los resúmenes y metodologías de todos los artículos
    contexto = ""
    for articulo in articulos:
        contexto += f"Resumen del artículo: {articulo.get('abstract', '')}\n\n"
        contexto += f"Metodología del artículo: {articulo.get('methodology', '')}\n\n"
        # Puedes incluir más campos si es necesario

    # Construir el prompt con el contexto necesario
    prompt = (
        "A partir de la siguiente información, genera la sección de 'Introducción' para un artículo científico sobre re-identificación de personas utilizando características biométricas y soft-biométricas. "
        "La introducción debe presentar el tema, justificar su importancia, revisar el estado del arte de manera concisa y establecer los objetivos del estudio. "
        "Debe incluir al menos 15 citas en formato IEEE, asegurando que los autores no se repitan. Utiliza referencias como marcadores de posición [1], [2], ..., [15], que luego serán reemplazadas por las referencias reales. "
        "Incluye la sección de referencias al final del documento.\n\n"
        f"{contexto}"
    )

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": "Eres un asistente especializado en ingeniería de aprendizaje automático y redacción científica en español."
            },
            {"role": "user", "content": prompt}
        ],
        max_tokens=2500  # Ajusta según sea necesario
    )

    introduccion = response['choices'][0]['message']['content'].strip()
    return introduccion

def main():
    # Rutas relativas al archivo JSON de entrada y a la carpeta de salida
    input_json_path = os.path.join('data', 'input', 'articulos_analizados.json')
    output_docx_folder = os.path.join('data', 'output')

    # Verificar si el archivo JSON existe
    if not os.path.exists(input_json_path):
        print(f"El archivo JSON no existe: {input_json_path}")
        return

    # Leer el JSON 'articulos_analizados'
    with open(input_json_path, 'r', encoding='utf-8') as json_file:
        articulos = json.load(json_file)

    # Generar la Introducción con citas
    introduccion = generar_introduccion_con_citas(articulos)

    # Crear el documento de Word y agregar el contenido
    document = Document()
    document.add_heading('Introducción', level=1)
    document.add_paragraph(introduccion)

    # Crear la carpeta de salida si no existe
    os.makedirs(output_docx_folder, exist_ok=True)
    output_docx_path = os.path.join(output_docx_folder, 'Introduccion.docx')

    document.save(output_docx_path)

    print(f"Documento 'Introduccion.docx' generado y guardado en {output_docx_path}.")

if __name__ == '__main__':
    main()

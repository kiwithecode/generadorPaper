import openai
import json
import os
from docx import Document

# Establece tu clave de API de OpenAI
openai.api_key = 'TU_CLAVE_DE_API_AQUÍ'  # Reemplaza con tu clave real

def generar_articulo(articulos):
    # Crear el contexto combinando los resúmenes, metodologías y resultados de todos los artículos
    contexto = ""
    for articulo in articulos:
        contexto += f"Resumen del artículo: {articulo.get('abstract', '')}\n\n"
        contexto += f"Metodología del artículo: {articulo.get('methodology', '')}\n\n"
        contexto += f"Resultados del artículo: {articulo.get('results', '')}\n\n"

    # Construir el prompt con el contexto necesario
    prompt = (
        "A partir de la siguiente información, genera un artículo científico completo sobre re-identificación de personas utilizando características biométricas y soft-biométricas. "
        "El artículo debe incluir las siguientes secciones: Resumen, Introducción, Materiales y Métodos, Resultados, Discusión y Conclusiones. "
        "Cada sección debe estar bien estructurada y redactada en un estilo científico y formal. "
        "Incluye definiciones de conceptos clave, ecuaciones relevantes y menciona las figuras correspondientes en los lugares adecuados. "
        "Sí incluye citas y referencias bibliográficas.\n\n"
        f"{contexto}"
    )

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": "Eres un asistente especializado en ingeniería de aprendizaje automático y redacción científica en español."
            },
            {"role": "user", "content": prompt}
        ],
        max_tokens=3500  # Ajusta según sea necesario
    )

    articulo = response['choices'][0]['message']['content'].strip()
    return articulo

def main():
    # Rutas relativas al archivo JSON de entrada y a la carpeta de salida
    input_json_path = os.path.join('data', 'input', 'articulos_analizados.json')
    output_docx_folder = os.path.join('data', 'output')

    # Verificar si el archivo JSON existe
    if not os.path.exists(input_json_path):
        print(f"El archivo JSON no existe: {input_json_path}")
        return

    # Leer el JSON 'articulos_analizados'
    with open(input_json_path, 'r', encoding='utf-8') as json_file:
        articulos = json.load(json_file)

    # Generar el artículo completo
    articulo = generar_articulo(articulos)

    # Crear el documento de Word y agregar el contenido
    document = Document()

    # Dividir el artículo en secciones basadas en encabezados
    secciones = articulo.split('\n\n')
    for seccion in secciones:
        linea = seccion.strip()
        if linea in ['Resumen', 'Introducción', 'Materiales y Métodos', 'Resultados', 'Discusión', 'Conclusiones']:
            document.add_heading(linea, level=1)
        else:
            document.add_paragraph(seccion)

    # Crear la carpeta de salida si no existe
    os.makedirs(output_docx_folder, exist_ok=True)
    output_docx_path = os.path.join(output_docx_folder, 'Articulo_Completo.docx')

    document.save(output_docx_path)

    print(f"Documento 'Articulo_Completo.docx' generado y guardado en {output_docx_path}.")

if __name__ == '__main__':
    main()

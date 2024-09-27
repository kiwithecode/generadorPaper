import openai
import json
import os
from docx import Document

# Establece tu clave de API de OpenAI
openai.api_key = 'TU_CLAVE_DE_API_AQUÍ'  # Reemplaza con tu clave real

def generar_materiales_metodos(articulos):
    # Crear el contexto combinando las metodologías de todos los artículos
    contexto = ""
    for articulo in articulos:
        contexto += f"Metodología del artículo: {articulo.get('methodology', '')}\n\n"

    # Construir el prompt con el contexto necesario
    prompt = (
        "En el contexto de la re-identificación de personas utilizando características biométricas y soft-biométricas, "
        "se requiere generar la sección de 'Materiales y Métodos' basada en el documento subido que contiene las secciones de "
        "'Materiales y Métodos' y 'Resultados'. Revisa y analiza y crea la sección solicitada, poniendo una introducción antes de empezar a describir esta sección. "
        "Dicha sección debe estar estructurada de manera técnica y detallada, redactada científicamente y con un hilo conductor claro. "
        "En ella se abordará la construcción del sistema, datasets de rostro (fijo), dataset de cuerpos (adquiridos en tiempo real), descriptores HoG y clasificadores SVM, "
        "identificación de personas, re-identificación de personas, métricas de evaluación, todos estos conceptos relacionados con la silueta de una persona. "
        "Esto se realizará en un circuito cerrado de cámaras no solapadas (4 cámaras), las cuales están estratégicamente ubicadas en dos esquinas, con dos cámaras por esquina orientadas en direcciones opuestas. "
        "Una cámara se encuentra posicionada para capturar los rostros de las personas que suben por las escaleras, mientras que la otra está configurada para monitorear la parte posterior, capturando imágenes de las personas que se alejan tras girar en la esquina. "
        "Esta disposición asegura una cobertura integral de la zona, permitiendo la captura de vistas frontales y traseras para un seguimiento completo. "
        "Se utilizará un dataset fijo de rostros con el cual se genera el modelo SVM de reconocimiento facial, mientras que los datasets de cuerpos se obtendrán en tiempo real a partir del circuito de cámaras. "
        "El sistema consta de dos procesos principales: identificación de personas y re-identificación de personas. Pon una introducción corta. "
        "Para el reconocimiento y localización del cuerpo y rostro de una persona, se emplea el algoritmo YOLO. "
        "En la identificación de personas, una vez localizado el rostro en tiempo real por una cámara (asociada a un procesador), se extraen las características faciales mediante el descriptor HoG y se identifican utilizando el modelo SVM previamente generado. "
        "Para el proceso de re-identificación de personas, una vez identificado el rostro con una de las cámaras, se procede a localizar el cuerpo, extrayendo las características de la silueta corporal utilizando HoG, y generando un modelo de silueta a través de SVM. "
        "Cuando se busca re-identificar a una persona sin que su rostro sea visible, se extraen las características de la silueta con HoG y estas se ingresan en el modelo SVM de silueta para re-identificar a la persona. "
        "Todos los procesos se realizan en tiempo real. En esta sección también se deben definir todos los conceptos de descriptores (técnicas de visión por computador, HoG para extraer las características del rostro y de la silueta) y clasificadores (modelos de machine learning, SVM para identificar el rostro y re-identificar a través de la silueta de una persona). "
        "Es obligatorio incluir las figuras correspondientes del documento subido en las partes más adecuadas de esta sección. "
        "Poner las ecuaciones de los descriptores, clasificadores y métricas de evaluación utilizadas. "
        "Además, se debe especificar por qué en este contexto no se utilizan métricas como la exactitud (accuracy) y el F1-score, y cuáles son las métricas más relevantes según lo descrito en el documento: mAP, ROC, Precision-Recall. "
        "Por favor, no pongas conclusión al final.\n\n"
        "Materiales y Métodos:\n\n"
        f"{contexto}"
    )

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": "Eres un asistente especializado en ingeniería de aprendizaje automático y redacción científica en español."
            },
            {"role": "user", "content": prompt}
        ],
        max_tokens=3000  # Ajusta según sea necesario
    )

    materiales_metodos = response['choices'][0]['message']['content'].strip()
    return materiales_metodos

def main():
    # Ruta relativa al archivo JSON de entrada
    input_json_path = os.path.join('data', 'input', 'articulos_analizados.json')
    output_docx_folder = os.path.join('data', 'output')

    # Verificar si el archivo JSON existe
    if not os.path.exists(input_json_path):
        print(f"El archivo JSON no existe: {input_json_path}")
        return

    # Leer el JSON 'articulos_analizados'
    with open(input_json_path, 'r', encoding='utf-8') as json_file:
        articulos = json.load(json_file)

    # Generar la sección de Materiales y Métodos
    materiales_metodos = generar_materiales_metodos(articulos)

    # Crear el documento de Word y agregar el contenido
    document = Document()
    document.add_heading('Materiales y Métodos', level=1)
    document.add_paragraph(materiales_metodos)

    # Crear la carpeta de salida si no existe
    os.makedirs(output_docx_folder, exist_ok=True)
    output_docx_path = os.path.join(output_docx_folder, 'Materiales_y_Metodos.docx')

    document.save(output_docx_path)

    print(f"Documento 'Materiales_y_Metodos.docx' generado y guardado en {output_docx_path}.")

if __name__ == '__main__':
    main()

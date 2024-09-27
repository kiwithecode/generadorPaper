import openai
import json
import os
from docx import Document

# Establece tu clave de API de OpenAI
openai.api_key = 'TU_CLAVE_DE_API_AQUÍ'  # Reemplaza con tu clave real

def generar_conclusiones(articulos):
    # Crear el contexto combinando los resultados y discusiones de todos los artículos
    contexto = ""
    for articulo in articulos:
        contexto += f"Resultados del artículo: {articulo.get('results', '')}\n\n"
        contexto += f"Discusión del artículo: {articulo.get('discussion', '')}\n\n"
        # Puedes incluir más campos si es necesario

    # Construir el prompt con el contexto necesario
    prompt = (
        "A partir de la siguiente información, genera la sección de 'Conclusiones' para un artículo científico sobre re-identificación de personas utilizando características biométricas y soft-biométricas. "
        "La sección de conclusiones debe resumir los hallazgos más importantes, discutir la relevancia de los resultados y sugerir posibles líneas de investigación futuras. "
        "Debe estar redactada en un estilo científico y formal. No incluyas citas ni referencias bibliográficas.\n\n"
        f"{contexto}"
    )

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": "Eres un asistente especializado en ingeniería de aprendizaje automático y redacción científica en español."
            },
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500  # Ajusta según sea necesario
    )

    conclusiones = response['choices'][0]['message']['content'].strip()
    return conclusiones

def main():
    # Rutas relativas al archivo JSON de entrada y a la carpeta de salida
    input_json_path = os.path.join('data', 'input', 'articulos_analizados.json')
    output_docx_folder = os.path.join('data', 'output')

    # Verificar si el archivo JSON existe
    if not os.path.exists(input_json_path):
        print(f"El archivo JSON no existe: {input_json_path}")
        return

    # Leer el JSON 'articulos_analizados'
    with open(input_json_path, 'r', encoding='utf-8') as json_file:
        articulos = json.load(json_file)

    # Generar la sección de Conclusiones
    conclusiones = generar_conclusiones(articulos)

    # Crear el documento de Word y agregar el contenido
    document = Document()
    document.add_heading('Conclusiones', level=1)
    document.add_paragraph(conclusiones)

    # Crear la carpeta de salida si no existe
    os.makedirs(output_docx_folder, exist_ok=True)
    output_docx_path = os.path.join(output_docx_folder, 'Conclusiones.docx')

    document.save(output_docx_path)

    print(f"Documento 'Conclusiones.docx' generado y guardado en {output_docx_path}.")

if __name__ == '__main__':
    main()
